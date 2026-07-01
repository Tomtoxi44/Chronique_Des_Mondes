"""
compress-docx.py
Post-traitement du docx généré par pandoc pour réduire le nombre de pages :
1. Supprime les paragraphes vides consécutifs
2. Réduit l'espacement des paragraphes normaux
3. Redimensionne les images à max 14 cm de large
4. Réduit la taille de police des blocs de code (Source Code → 8pt)
"""

from docx import Document
from docx.shared import Cm, Pt, Emu
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
import os

INPUT  = r"D:\Projet_Perso\Chronique_Des_Mondes\dossier-bloc2-v2.docx"
OUTPUT = r"D:\Projet_Perso\Chronique_Des_Mondes\dossier-bloc2.docx"

MAX_IMG_WIDTH = Cm(14)   # largeur max image en EMU

doc = Document(INPUT)

# ── 1. Ajuster marges de page ─────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(1.8)
    section.right_margin  = Cm(1.8)

# ── 2. Supprimer paragraphes vides consécutifs ────────────────────────────
prev_empty = False
to_remove = []
for para in doc.paragraphs:
    text = para.text.strip()
    style_name = para.style.name if para.style else ''
    is_empty = (text == '' and 'Heading' not in style_name)
    if is_empty and prev_empty:
        to_remove.append(para)
    prev_empty = is_empty

removed = 0
for para in to_remove:
    try:
        p = para._element
        p.getparent().remove(p)
        removed += 1
    except Exception:
        pass
print(f"Paragraphes vides supprimés : {removed}")

# ── 3. Compresser espacement des paragraphes normaux ──────────────────────
compressed = 0
for para in doc.paragraphs:
    style_name = para.style.name if para.style else ''
    if style_name in ('Normal', 'Body Text', '') or style_name.startswith('List'):
        pf = para.paragraph_format
        if pf.space_before is None or pf.space_before > Pt(2):
            pf.space_before = Pt(0)
        if pf.space_after is None or pf.space_after > Pt(4):
            pf.space_after = Pt(3)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        compressed += 1

print(f"Paragraphes compressés : {compressed}")

# ── 4. Réduire taille police des blocs de code ───────────────────────────
code_paras = 0
for para in doc.paragraphs:
    style_name = para.style.name if para.style else ''
    if 'Source Code' in style_name or 'Verbatim' in style_name or 'Code' in style_name:
        for run in para.runs:
            if run.font.size is None or run.font.size > Pt(8.5):
                run.font.size = Pt(8)
        pf = para.paragraph_format
        pf.space_before = Pt(1)
        pf.space_after  = Pt(1)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        code_paras += 1

print(f"Paragraphes de code ajustés : {code_paras}")

# ── 5. Redimensionner images trop larges ─────────────────────────────────
resized = 0
for para in doc.paragraphs:
    for run in para.runs:
        for drawing in run._r.findall('.//' + qn('a:blip'), run._r.nsmap):
            pass
        # Recherche inline images
        for inline in run._r.findall('.//' + qn('wp:inline')):
            extent = inline.find(qn('wp:extent'))
            if extent is not None:
                cx = int(extent.get('cx', 0))
                if cx > MAX_IMG_WIDTH:
                    ratio = MAX_IMG_WIDTH / cx
                    cy = int(extent.get('cy', 0))
                    extent.set('cx', str(int(MAX_IMG_WIDTH)))
                    extent.set('cy', str(int(cy * ratio)))
                    resized += 1

print(f"Images redimensionnées : {resized}")

# ── 6. Sauvegarder ────────────────────────────────────────────────────────
doc.save(OUTPUT)
size_ko = os.path.getsize(OUTPUT) // 1024
print(f"\nFichier final : {OUTPUT}")
print(f"Taille        : {size_ko} Ko")
