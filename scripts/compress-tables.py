"""
compress-tables.py
Compresse les tableaux du docx : police 9pt, espacement réduit, pas de ligne vide dans les cellules.
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING
import os

INPUT  = r"D:\Projet_Perso\Chronique_Des_Mondes\dossier-bloc2.docx"
OUTPUT = r"D:\Projet_Perso\Chronique_Des_Mondes\dossier-bloc2.docx"

doc = Document(INPUT)

table_cells = 0
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                # Police 9pt dans les cellules
                for run in para.runs:
                    if run.font.size is None or run.font.size > Pt(9.5):
                        run.font.size = Pt(9)
                # Espacement minimal
                pf = para.paragraph_format
                pf.space_before      = Pt(1)
                pf.space_after       = Pt(1)
                pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
                table_cells += 1

print(f"Cellules de tableau traitées : {table_cells} dans {len(doc.tables)} tables")

doc.save(OUTPUT)
print(f"Sauvegardé : {OUTPUT} ({os.path.getsize(OUTPUT)//1024} Ko)")
