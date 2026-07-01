"""
build-alt-docx.py
Crée dossier-bloc2-alt.docx :
- Reprend dossier-bloc2.docx comme base
- Remplace l'image ER (diagram_1.png) par un tableau Word natif propre
- Chaque entité = bloc titre + tableau champs/type/description
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, copy

INPUT  = r"D:\Projet_Perso\Chronique_Des_Mondes\dossier-bloc2.docx"
OUTPUT = r"D:\Projet_Perso\Chronique_Des_Mondes\dossier-bloc2-alt.docx"

# ── Données ER réelles (modèles EF Core) ─────────────────────────────────
ENTITIES = [
    {
        "name": "USERS",
        "color": RGBColor(0x1E, 0x1B, 0x4B),
        "fields": [
            ("Id",           "int",       "PK – Identifiant unique"),
            ("Email",        "string",    "UK – Adresse e-mail (unique)"),
            ("PasswordHash", "string",    "Hash BCrypt du mot de passe"),
            ("Role",         "string",    "Rôle applicatif (User / Admin)"),
            ("Nickname",     "string",    "Pseudo affiché"),
            ("CreatedAt",    "datetime",  "Date de création du compte"),
        ]
    },
    {
        "name": "WORLDS",
        "color": RGBColor(0x31, 0x2E, 0x81),
        "fields": [
            ("Id",           "int",       "PK"),
            ("Name",         "string",    "Nom du monde"),
            ("Description",  "string",    "Description"),
            ("GameType",     "enum",      "Type de jeu (DnD5e, Generic…)"),
            ("CreatedBy",    "int FK",    "→ USERS.Id (créateur / MJ)"),
            ("InviteToken",  "string?",   "Token d'invitation (nullable)"),
            ("CreatedAt",    "datetime",  "Date de création"),
        ]
    },
    {
        "name": "CAMPAIGNS",
        "color": RGBColor(0x4F, 0x46, 0xE5),
        "fields": [
            ("Id",           "int",       "PK"),
            ("Name",         "string",    "Nom de la campagne"),
            ("Description",  "string?",   "Description"),
            ("Status",       "enum",      "Planning | Active | OnHold | Completed"),
            ("WorldId",      "int FK",    "→ WORLDS.Id"),
            ("CreatedBy",    "int FK",    "→ USERS.Id (créateur MJ)"),
            ("CreatedAt",    "datetime",  "Date de création"),
        ]
    },
    {
        "name": "CHARACTERS",
        "color": RGBColor(0x7C, 0x3A, 0xED),
        "fields": [
            ("Id",           "int",       "PK"),
            ("Name",         "string",    "Nom du personnage"),
            ("GameType",     "enum",      "Type de jeu"),
            ("UserId",       "int FK",    "→ USERS.Id (propriétaire)"),
            ("IsLocked",     "bool",      "Verrouillé par une session active"),
            ("CreatedAt",    "datetime",  "Date de création"),
        ]
    },
    {
        "name": "WORLD_CHARACTERS",
        "color": RGBColor(0xA7, 0x8B, 0xFA),
        "fields": [
            ("WorldId",      "int FK",    "PK + → WORLDS.Id"),
            ("CharacterId",  "int FK",    "PK + → CHARACTERS.Id"),
            ("JoinedAt",     "datetime",  "Date d'entrée dans le monde"),
            ("ProfileJson",  "nvarchar?", "Stats spécifiques au système de jeu"),
        ]
    },
    {
        "name": "SESSIONS",
        "color": RGBColor(0x0E, 0x7C, 0x86),
        "fields": [
            ("Id",           "int",       "PK"),
            ("Status",       "enum",      "Pending | Active | Ended"),
            ("CampaignId",   "int FK",    "→ CAMPAIGNS.Id"),
            ("StartedById",  "int FK",    "→ USERS.Id (MJ qui a lancé)"),
            ("StartedAt",    "datetime?", "Date de démarrage"),
            ("EndedAt",      "datetime?", "Date de fin"),
        ]
    },
    {
        "name": "COMBAT_ACTIONS",
        "color": RGBColor(0xBE, 0x12, 0x3C),
        "fields": [
            ("Id",           "int",       "PK"),
            ("DiceExpr",     "string",    "Expression de dés (ex: 2d6+3)"),
            ("RollResult",   "int",       "Résultat RNG côté serveur"),
            ("ActionType",   "string",    "Type d'action (Attack, Skill…)"),
            ("SessionId",    "int FK",    "→ SESSIONS.Id"),
            ("CharacterId",  "int FK",    "→ CHARACTERS.Id"),
            ("RolledAt",     "datetime",  "Timestamp de l'action"),
        ]
    },
]

RELATIONS = [
    "USERS        ||--o{ WORLDS           : crée (CreatedBy)",
    "WORLDS       ||--o{ CAMPAIGNS        : héberge",
    "USERS        ||--o{ CHARACTERS       : possède (UserId)",
    "WORLDS       }o--o{ CHARACTERS       : via WORLD_CHARACTERS",
    "CAMPAIGNS    ||--o{ SESSIONS         : héberge",
    "SESSIONS     ||--o{ COMBAT_ACTIONS   : enregistre",
    "CHARACTERS   ||--o{ COMBAT_ACTIONS   : réalise",
]

# ── Helpers ───────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: RGBColor):
    """Définit la couleur de fond d'une cellule."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    r = f"{hex_color[0]:02X}"
    g = f"{hex_color[1]:02X}"
    b = f"{hex_color[2]:02X}"
    shd.set(qn('w:fill'), f"{r}{g}{b}")
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

def set_cell_border(table):
    """Active les bordures simples sur toutes les cellules du tableau."""
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                border = OxmlElement(f'w:{side}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:color'), 'C7D2FE')
                tcBorders.append(border)
            tcPr.append(tcBorders)


def add_er_table(doc, entity):
    """Ajoute un tableau entité au document."""
    # En-tête entité
    p = doc.add_paragraph()
    run = p.add_run(f"  {entity['name']}")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(0)

    # Couleur de fond du paragraphe-titre
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    c = entity['color']
    shd.set(qn('w:fill'), f"{c[0]:02X}{c[1]:02X}{c[2]:02X}")
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)

    # Tableau champs
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    # En-tête colonnes
    hdr = table.rows[0].cells
    for i, txt in enumerate(["Champ", "Type", "Description"]):
        hdr[i].text = txt
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_bg(hdr[i], RGBColor(0xE0, 0xE7, 0xFF))

    # Lignes de champs
    for field_name, field_type, field_desc in entity['fields']:
        row = table.add_row()
        row.cells[0].text = field_name
        row.cells[1].text = field_type
        row.cells[2].text = field_desc
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(9)

    set_cell_border(table)

    # Largeurs des colonnes
    for i, width in enumerate([Cm(3.5), Cm(2.5), Cm(10)]):
        for row in table.rows:
            row.cells[i].width = width

    doc.add_paragraph()


# ── Localiser et remplacer le diagramme ER dans le docx ──────────────────

doc = Document(INPUT)

# Trouver le paragraphe contenant diagram_1.png (ER)
er_para_idx = None
for i, para in enumerate(doc.paragraphs):
    # Les images sont dans les runs via des éléments drawing
    for run in para.runs:
        for drawing in run._r.findall('.//' + qn('wp:inline')):
            # Vérifier si c'est le 2e diagramme (ER = diagram_1)
            blips = run._r.findall('.//' + qn('a:blip'))
            for blip in blips:
                rEmbed = blip.get(qn('r:embed'))
                if rEmbed:
                    # Chercher dans les relations du document
                    pass
        # Alternative : chercher par position approximative
        xml = run._r.xml if hasattr(run._r, 'xml') else ''
        if 'drawing' in xml or 'blip' in xml:
            er_para_idx = i  # on prendra le 2e
            break

# Approche plus simple : trouver tous les paragraphes avec images
image_paras = []
for i, para in enumerate(doc.paragraphs):
    for run in para.runs:
        drawings = run._r.findall('.//' + qn('wp:inline'))
        if drawings:
            image_paras.append(i)
            break

print(f"Paragraphes avec images trouvés : {image_paras}")

# Le 2e diagramme est l'ER (diagram_1)
if len(image_paras) >= 2:
    er_para_idx = image_paras[1]
    er_para = doc.paragraphs[er_para_idx]
    parent = er_para._element.getparent()
    er_pos = list(parent).index(er_para._element)

    # Supprimer l'image ER
    parent.remove(er_para._element)
    print(f"Image ER supprimée (position {er_para_idx})")

    # Insérer un marqueur temporaire pour récupérer la position
    marker = doc.add_paragraph("__ER_PLACEHOLDER__")
    marker_el = marker._element
    parent.insert(er_pos, marker_el)
    doc.paragraphs[-1]._element.getparent().remove(doc.paragraphs[-1]._element)

    # Insérer les tables ER à la place du marqueur
    # On ne peut pas insérer facilement à une position arbitraire avec python-docx
    # Stratégie : sauvegarder en intermédiaire, reconstruire

# Sauvegarder version intermédiaire et reconstruire
doc.save(OUTPUT)

# Reconstruire proprement avec les tables ER
doc2 = Document(OUTPUT)
for i, para in enumerate(doc2.paragraphs):
    if '__ER_PLACEHOLDER__' in para.text:
        placeholder = para._element
        placeholder_parent = placeholder.getparent()
        placeholder_pos = list(placeholder_parent).index(placeholder)

        # Créer un doc temporaire pour générer les éléments ER
        tmp_doc = Document()
        for section in tmp_doc.sections:
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)
            section.left_margin = Cm(1.8)
            section.right_margin = Cm(1.8)

        # Titre de la section ER
        h = tmp_doc.add_paragraph("Schéma Entité-Relation — Modèles EF Core")
        h.style = tmp_doc.styles['Normal']
        h.runs[0].bold = True
        h.runs[0].font.size = Pt(11)

        # Relations
        rel_para = tmp_doc.add_paragraph()
        rel_para.add_run("Relations : ").bold = True
        for rel in RELATIONS:
            tmp_doc.add_paragraph(f"• {rel}", style='Normal').runs[0].font.size = Pt(9)

        tmp_doc.add_paragraph()

        # Entités
        for entity in ENTITIES:
            add_er_table(tmp_doc, entity)

        tmp_doc.save(r"D:\Projet_Perso\Chronique_Des_Mondes\docs\er-block.docx")

        # Insérer les éléments du tmp_doc dans le doc2
        tmp = Document(r"D:\Projet_Perso\Chronique_Des_Mondes\docs\er-block.docx")
        elements_to_insert = list(tmp.element.body)[:-1]  # Exclure sectPr

        # Copier les namespace declarations si nécessaire
        for j, el in enumerate(reversed(elements_to_insert)):
            el_copy = copy.deepcopy(el)
            placeholder_parent.insert(placeholder_pos, el_copy)

        # Supprimer le placeholder
        placeholder_parent.remove(placeholder)
        print(f"Tables ER insérées ({len(elements_to_insert)} éléments)")
        break

doc2.save(OUTPUT)
print(f"\nFichier alternatif : {OUTPUT}")
print(f"Taille : {os.path.getsize(OUTPUT)//1024} Ko")

