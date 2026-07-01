"""
resize-text.py
Augmente la taille du texte dans le docx :
- Corps de texte (Normal) : 11pt
- Tableaux : 10pt
- Blocs de code : 9pt (lisible mais compact)
- Titres : légèrement augmentés
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
import os

PATH = r"D:\Projet_Perso\Chronique_Des_Mondes\dossier-bloc2.docx"

doc = Document(PATH)

# ── Corps de texte ────────────────────────────────────────────────────────
normal_count = 0
for para in doc.paragraphs:
    style_name = para.style.name if para.style else ''

    # Blocs de code → 9pt
    if 'Source Code' in style_name or 'Verbatim' in style_name or 'Code' in style_name:
        for run in para.runs:
            run.font.size = Pt(9)
        continue

    # Titres → ne pas toucher (ils ont leurs propres tailles)
    if 'Heading' in style_name:
        continue

    # Tout le reste (Normal, Body, listes, etc.) → 11pt
    for run in para.runs:
        run.font.size = Pt(11)
    normal_count += 1

print(f"Paragraphes texte mis à 11pt : {normal_count}")

# ── Tableaux → 10pt ───────────────────────────────────────────────────────
table_runs = 0
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    table_runs += 1

print(f"Runs de tableau mis à 10pt : {table_runs}")

doc.save(PATH)
print(f"\nSauvegardé : {PATH} ({os.path.getsize(PATH)//1024} Ko)")
